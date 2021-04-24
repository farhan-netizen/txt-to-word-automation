# -*- coding: utf-8 -*-
"""
Created on Wed Mar 31 17:59:20 2021

@author: Farhan Ashraf
"""
from docx import Document
import re
import os

import xlrd

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt

path_of_excel_path = 'C:/Users/Farhan Ashraf/Desktop/holderID.xlsx'
excel_having_holder_id = xlrd.open_workbook(path_of_excel_path)
sheet = excel_having_holder_id.sheet_by_index(0)
holder_id = int(sheet.cell_value(1,0))

path_of_txt_file = 'C:/Users/Farhan Ashraf/Desktop/'
files = os.listdir(path_of_txt_file)

document = Document()

def change_orientation():
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    

    return new_section

for i in files:
    if i == 'operations' + str(holder_id) + '.txt':
        print(i[:-4])
        
        document.add_heading(i, 0)
        txt_file_content = open('C:/Users/farhan Ashraf/Desktop' +  i ).readlines()
        count = 0
        
        for k in txt_file_content:
            count += 1
            
        print(count)
        
        try:
            paragraph = document.add_paragraph(txt_file_content[0])
            paragraph.add_run().add_break()
            
            for j in range(1, count):
                c = txt_file_content[j].split("\t")
                balance = int(float(c[4]))
                paragraph.add_run(txt_file_content[j][:-1])
                if (c[3][0] + c[3][1] + c[3][2] + c[3][3]) == 'UTNA':
                    if c[1] in ["NP"]:
                        print("Need to be checked")
            
                        run = paragraph.add_run('\tCheck in app')
                        run.font.color.rgb = red
                    elif balance < 100:
                        run = paragraph.add_run('\tU/T\n')
                        run.font.color.rgb = red
                        
                    elif balance > 100: 
                        run = paragraph.add_run('\tN/A\n')
                        run.font.color.rgb = red
                        
                elif balance < 100:
                    run = paragraph.add_run('\tU/T\n')
                    run.font.color.rgb = red
                        
                elif balance > 100:
                    run = paragraph.add_run('\tN/A\n')
                    run.font.color.rgb = red
        except:
            pass

        font = document.styles['Normal'].font
        font.name = 'Courier New'
        font.size = Pt(8)
        document.save('C:/Users/Farhan Ashraf/Desktop/' + i[:-4] + '.docx')    


